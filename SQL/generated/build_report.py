from __future__ import annotations

import os
import re
import shutil
import subprocess
import textwrap
from pathlib import Path

import pypdfium2 as pdfium
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
SQL_DIR = ROOT / "SQL"
OUT_DIR = SQL_DIR / "generated"
STUDENT_ID = os.environ.get("SQL_REPORT_STUDENT_ID", "<学号>")
STUDENT_NAME = os.environ.get("SQL_REPORT_STUDENT_NAME", "<姓名>")
REPORT_PATH = SQL_DIR / "《数据库系统原理》课程实验报告模板2_模板.docx"
PDF_PATH = OUT_DIR / "《数据库系统原理》课程实验报告模板2_模板.pdf"
RENDER_DIR = OUT_DIR / "rendered_pages"

CONTAINER = "sql2022-lab"
SA_PASSWORD = "CodexLab!2026"
SQLCMD = "/opt/mssql-tools18/bin/sqlcmd"


SETUP_SQL = r"""
USE master;
IF DB_ID(N'GradeManager') IS NOT NULL
BEGIN
    ALTER DATABASE GradeManager SET SINGLE_USER WITH ROLLBACK IMMEDIATE;
    DROP DATABASE GradeManager;
END;
GO

CREATE DATABASE GradeManager COLLATE Chinese_PRC_CI_AS;
GO

USE GradeManager;
GO

CREATE TABLE Class (
    Clno CHAR(5) NOT NULL PRIMARY KEY,
    Speciality VARCHAR(20) NOT NULL,
    Inyear CHAR(4) NOT NULL,
    Number INTEGER NULL,
    Monitor CHAR(7) NULL
);

CREATE TABLE Student (
    Sno CHAR(7) NOT NULL PRIMARY KEY,
    Sname VARCHAR(20) NOT NULL,
    Ssex CHAR(2) NOT NULL,
    Sbirth DATE NULL,
    Clno CHAR(5) NOT NULL
);

CREATE TABLE Course (
    Cno CHAR(3) NOT NULL PRIMARY KEY,
    Cname VARCHAR(20) NOT NULL,
    Ccredit SMALLINT NULL,
    Cpno CHAR(3) NULL
);

CREATE TABLE Grade (
    Sno CHAR(7) NOT NULL,
    Cno CHAR(3) NOT NULL,
    Gmark NUMERIC(4,1) NULL,
    CONSTRAINT PK_Grade PRIMARY KEY (Sno, Cno)
);

INSERT INTO Student (Sno, Sname, Ssex, Sbirth, Clno) VALUES
('2020101', N'李勇', N'男', '2002-08-09', '20311'),
('2020102', N'刘诗晨', N'女', '2003-04-01', '20311'),
('2020103', N'王一鸣', N'男', '2002-12-25', '20312'),
('2020104', N'张婷婷', N'女', '2002-10-01', '20312'),
('2021101', N'李勇敏', N'女', '2003-11-11', '21311'),
('2021102', N'贾向东', N'男', '2003-12-12', '21311'),
('2021103', N'陈宝玉', N'男', '2004-05-01', '21311'),
('2021104', N'张逸凡', N'男', '2005-01-01', '21311');

INSERT INTO Course (Cno, Cname, Ccredit, Cpno) VALUES
('1', N'数据库系统原理', 4, '5'),
('2', N'计算机系统结构', 3, '8'),
('3', N'数字电路设计', 2, NULL),
('4', N'操作系统', 4, '8'),
('5', N'数据结构', 4, '7'),
('6', N'软件工程', 2, '1'),
('7', N'C语言', 4, NULL),
('8', N'计算机组成原理', 4, '3');

INSERT INTO Class (Clno, Speciality, Inyear, Number, Monitor) VALUES
('20311', N'软件工程', '2020', 35, '2020101'),
('20312', N'计算机科学与技术', '2020', 38, '2020103'),
('21311', N'软件工程', '2021', 40, '2021103');

INSERT INTO Grade (Sno, Cno, Gmark) VALUES
('2020101', '1', 92), ('2020101', '3', 88), ('2020101', '5', 86),
('2020102', '1', 78), ('2020102', '6', 55),
('2020103', '3', 65), ('2020103', '6', 78), ('2020103', '5', 66),
('2020104', '1', 54), ('2020104', '6', 83),
('2021101', '2', 70), ('2021101', '4', 65),
('2021102', '2', 80), ('2021102', '4', 90),
('2020103', '1', 83), ('2020103', '2', 76), ('2020103', '4', 56), ('2020103', '7', 88);

CREATE TABLE Employee (
    Eno CHAR(6) NOT NULL PRIMARY KEY,
    Ename VARCHAR(20) NOT NULL,
    Title VARCHAR(20) NOT NULL
);

CREATE TABLE Salary (
    Eno CHAR(6) NOT NULL PRIMARY KEY,
    Basepay DECIMAL(10,2) NOT NULL,
    Service DECIMAL(10,2) NOT NULL
);

INSERT INTO Employee (Eno, Ename, Title) VALUES
('E001', N'赵强', N'工程师'),
('E002', N'钱华', N'助理工程师'),
('E003', N'孙敏', N'工程师'),
('E004', N'周宁', N'经理');

INSERT INTO Salary (Eno, Basepay, Service) VALUES
('E001', 3600, 420),
('E002', 2800, 300),
('E003', 4100, 500),
('E004', 5200, 800);
GO

CREATE OR ALTER FUNCTION dbo.AgeOf(@birth DATE)
RETURNS INT
AS
BEGIN
    RETURN DATEDIFF(YEAR, @birth, CAST(GETDATE() AS DATE))
           - CASE
               WHEN DATEADD(YEAR, DATEDIFF(YEAR, @birth, CAST(GETDATE() AS DATE)), @birth) > CAST(GETDATE() AS DATE)
               THEN 1 ELSE 0
             END;
END;
GO

SELECT N'GradeManager 初始化完成' AS Info,
       (SELECT COUNT(*) FROM Student) AS StudentRows,
       (SELECT COUNT(*) FROM Course) AS CourseRows,
       (SELECT COUNT(*) FROM Class) AS ClassRows,
       (SELECT COUNT(*) FROM Grade) AS GradeRows;
GO
"""


EXPERIMENT5_SQL = r"""
USE GradeManager;
GO

PRINT N'【实验五】SELECT 语句高级格式和完整格式的使用';

SELECT N'13(1) 与李勇在同一个班级的学生信息' AS Item;
SELECT Sno, Sname, Ssex, Sbirth, Clno
FROM Student
WHERE Clno = (SELECT Clno FROM Student WHERE Sname = N'李勇')
  AND Sname <> N'李勇';

SELECT N'13(2) 与李勇有相同选修课程的学生信息' AS Item;
SELECT DISTINCT s.Sno, s.Sname, s.Ssex, s.Sbirth, s.Clno
FROM Student AS s
WHERE s.Sno <> (SELECT Sno FROM Student WHERE Sname = N'李勇')
  AND s.Sno IN (
      SELECT g.Sno
      FROM Grade AS g
      WHERE g.Cno IN (
          SELECT Cno
          FROM Grade
          WHERE Sno = (SELECT Sno FROM Student WHERE Sname = N'李勇')
      )
  )
ORDER BY s.Sno;

SELECT N'13(3) 年龄介于李勇和25岁之间的学生信息' AS Item;
SELECT Sno, Sname, dbo.AgeOf(Sbirth) AS Age, Sbirth, Clno
FROM Student
WHERE Sname <> N'李勇'
  AND dbo.AgeOf(Sbirth) BETWEEN
      (SELECT dbo.AgeOf(Sbirth) FROM Student WHERE Sname = N'李勇') AND 25
ORDER BY Age, Sno;

SELECT N'13(4) 选修了操作系统的学生学号和姓名' AS Item;
SELECT s.Sno, s.Sname
FROM Student AS s
WHERE s.Sno IN (
    SELECT g.Sno
    FROM Grade AS g
    WHERE g.Cno = (SELECT Cno FROM Course WHERE Cname = N'操作系统')
)
ORDER BY s.Sno;

SELECT N'13(5) 没有选修1号课程的学生姓名' AS Item;
SELECT s.Sname
FROM Student AS s
WHERE NOT EXISTS (
    SELECT 1
    FROM Grade AS g
    WHERE g.Sno = s.Sno AND g.Cno = '1'
)
ORDER BY s.Sno;

SELECT N'13(6) 每个学生超过本人平均成绩的学号和课程号' AS Item;
SELECT g.Sno, g.Cno, g.Gmark
FROM Grade AS g
WHERE g.Gmark > (
    SELECT AVG(g2.Gmark)
    FROM Grade AS g2
    WHERE g2.Sno = g.Sno
)
ORDER BY g.Sno, g.Cno;

SELECT N'13(7) 选修了全部课程的学生姓名' AS Item;
SELECT s.Sname
FROM Student AS s
WHERE NOT EXISTS (
    SELECT 1
    FROM Course AS c
    WHERE NOT EXISTS (
        SELECT 1
        FROM Grade AS g
        WHERE g.Sno = s.Sno AND g.Cno = c.Cno
    )
);

SELECT N'13(8) 数据库系统原理成绩高于该课程平均分的学生学号、姓名、成绩' AS Item;
SELECT s.Sno, s.Sname, g.Gmark
FROM Student AS s
JOIN Grade AS g ON s.Sno = g.Sno
JOIN Course AS c ON g.Cno = c.Cno
WHERE c.Cname = N'数据库系统原理'
  AND g.Gmark > (
      SELECT AVG(g2.Gmark)
      FROM Grade AS g2
      JOIN Course AS c2 ON g2.Cno = c2.Cno
      WHERE c2.Cname = N'数据库系统原理'
  )
ORDER BY g.Gmark DESC;

SELECT N'13(9) 每个班中数据库系统原理成绩高于本班平均分的学生' AS Item;
SELECT s.Clno, s.Sno, s.Sname, g.Gmark
FROM Student AS s
JOIN Grade AS g ON s.Sno = g.Sno
JOIN Course AS c ON g.Cno = c.Cno
WHERE c.Cname = N'数据库系统原理'
  AND g.Gmark > (
      SELECT AVG(g2.Gmark)
      FROM Student AS s2
      JOIN Grade AS g2 ON s2.Sno = g2.Sno
      JOIN Course AS c2 ON g2.Cno = c2.Cno
      WHERE c2.Cname = N'数据库系统原理'
        AND s2.Clno = s.Clno
  )
ORDER BY s.Clno, g.Gmark DESC;

SELECT N'13(10) 至少选修了2020101号学生全部课程的学生学号' AS Item;
SELECT s.Sno
FROM Student AS s
WHERE NOT EXISTS (
    SELECT 1
    FROM Grade AS gy
    WHERE gy.Sno = '2020101'
      AND NOT EXISTS (
          SELECT 1
          FROM Grade AS gx
          WHERE gx.Sno = s.Sno AND gx.Cno = gy.Cno
      )
)
ORDER BY s.Sno;

SELECT N'14(1) 选修3号课程的学生学号及成绩，按成绩降序' AS Item;
SELECT Sno, Gmark
FROM Grade
WHERE Cno = '3'
ORDER BY Gmark DESC;

SELECT N'14(2) 全体学生信息，按班级升序、年龄降序' AS Item;
SELECT Sno, Sname, Ssex, dbo.AgeOf(Sbirth) AS Age, Clno
FROM Student
ORDER BY Clno ASC, Age DESC;

SELECT N'14(3) 每个课程号及相应选课人数' AS Item;
SELECT Cno, COUNT(*) AS StudentCount
FROM Grade
GROUP BY Cno
ORDER BY Cno;

SELECT N'14(4) 选修三门以上课程的学生学号' AS Item;
SELECT Sno, COUNT(*) AS CourseCount
FROM Grade
GROUP BY Sno
HAVING COUNT(*) >= 3
ORDER BY Sno;

SELECT N'14(5) 至少选修1号和2号课程的学生学号和姓名' AS Item;
SELECT s.Sno, s.Sname
FROM Student AS s
WHERE s.Sno IN (
    SELECT Sno
    FROM Grade
    WHERE Cno IN ('1', '2')
    GROUP BY Sno
    HAVING COUNT(DISTINCT Cno) = 2
)
ORDER BY s.Sno;

SELECT N'14(6) 每门课程成绩前三名的学生学号、姓名、课程号和成绩' AS Item;
WITH RankedGrade AS (
    SELECT g.Sno, s.Sname, g.Cno, g.Gmark,
           ROW_NUMBER() OVER (PARTITION BY g.Cno ORDER BY g.Gmark DESC, g.Sno) AS rn
    FROM Grade AS g
    JOIN Student AS s ON s.Sno = g.Sno
)
SELECT Sno, Sname, Cno, Gmark
FROM RankedGrade
WHERE rn <= 3
ORDER BY Cno, Gmark DESC;

SELECT N'14(7) 每个学生的总学分，按总学分降序' AS Item;
SELECT s.Sno, s.Sname, SUM(c.Ccredit) AS TotalCredit
FROM Student AS s
JOIN Grade AS g ON s.Sno = g.Sno
JOIN Course AS c ON g.Cno = c.Cno
GROUP BY s.Sno, s.Sname
ORDER BY TotalCredit DESC, s.Sno;
GO
"""


EXPERIMENT6_SQL = r"""
USE GradeManager;
GO

PRINT N'【实验六】SQL 的存储操作';

BEGIN TRAN;

SELECT N'15(1) 将20311班全体学生的成绩置0值（事务内演示）' AS Item;
UPDATE g
SET Gmark = 0
FROM Grade AS g
JOIN Student AS s ON s.Sno = g.Sno
WHERE s.Clno = '20311';
SELECT @@ROWCOUNT AS UpdatedRows;
SELECT s.Clno, g.Sno, g.Cno, g.Gmark
FROM Grade AS g
JOIN Student AS s ON s.Sno = g.Sno
WHERE s.Clno = '20311'
ORDER BY g.Sno, g.Cno;

ROLLBACK;

BEGIN TRAN;

SELECT N'15(2) 删除2021级软件工程学生的选课记录（事务内演示）' AS Item;
DELETE g
FROM Grade AS g
WHERE g.Sno IN (
    SELECT s.Sno
    FROM Student AS s
    JOIN Class AS c ON s.Clno = c.Clno
    WHERE c.Inyear = '2021' AND c.Speciality LIKE N'%软件%'
);
SELECT @@ROWCOUNT AS DeletedRows;

ROLLBACK;

BEGIN TRAN;

SELECT N'15(3) 学生李勇退学，删除与他有关的记录（事务内演示）' AS Item;
DELETE FROM Grade
WHERE Sno = (SELECT Sno FROM Student WHERE Sname = N'李勇');
SELECT @@ROWCOUNT AS DeletedGradeRows;
DELETE FROM Student
WHERE Sname = N'李勇';
SELECT @@ROWCOUNT AS DeletedStudentRows;

ROLLBACK;

SELECT N'15(4) 对每个班求学生平均年龄，并把结果存入数据库' AS Item;
DROP TABLE IF EXISTS ClassAvgAge;
SELECT c.Clno, c.Speciality, c.Inyear,
       CAST(AVG(CAST(dbo.AgeOf(s.Sbirth) AS DECIMAL(6,2))) AS DECIMAL(6,2)) AS AvgAge
INTO ClassAvgAge
FROM Class AS c
JOIN Student AS s ON c.Clno = s.Clno
GROUP BY c.Clno, c.Speciality, c.Inyear;
SELECT * FROM ClassAvgAge ORDER BY Clno;

SELECT N'验证工程师基本工资增加100的UPDATE语句（事务内演示）' AS Item;
BEGIN TRAN;
SELECT N'更新前' AS Stage, e.Eno, e.Ename, e.Title, s.Basepay
FROM Employee AS e
JOIN Salary AS s ON e.Eno = s.Eno
WHERE e.Title = N'工程师'
ORDER BY e.Eno;

UPDATE Salary
SET Basepay = Basepay + 100
WHERE Eno IN (
    SELECT Eno
    FROM Employee
    WHERE Title = N'工程师'
);
SELECT @@ROWCOUNT AS UpdatedEngineerRows;

SELECT N'更新后（事务内）' AS Stage, e.Eno, e.Ename, e.Title, s.Basepay
FROM Employee AS e
JOIN Salary AS s ON e.Eno = s.Eno
WHERE e.Title = N'工程师'
ORDER BY e.Eno;
ROLLBACK;
GO
"""


EXPERIMENT7_SQL = r"""
USE GradeManager;
GO

PRINT N'【实验七】视图的建立及操作';
GO

CREATE OR ALTER VIEW Stu_20312_1
AS
SELECT s.Sno, s.Sname, s.Ssex, s.Sbirth, s.Clno, g.Cno, g.Gmark
FROM Student AS s
JOIN Grade AS g ON s.Sno = g.Sno
WHERE s.Clno = '20312' AND g.Cno = '1';
GO

CREATE OR ALTER VIEW Stu_20312_2
AS
SELECT s.Sno, s.Sname, s.Ssex, s.Sbirth, s.Clno, g.Cno, g.Gmark
FROM Student AS s
JOIN Grade AS g ON s.Sno = g.Sno
WHERE s.Clno = '20312' AND g.Cno = '1' AND g.Gmark < 60;
GO

CREATE OR ALTER VIEW Stu_age
AS
SELECT Sno, Sname, dbo.AgeOf(Sbirth) AS Age
FROM Student;
GO

SELECT N'16(1) 20312班选修1号课程的学生视图 Stu_20312_1' AS Item;
SELECT Sno, Sname, Cno, Gmark FROM Stu_20312_1 ORDER BY Sno;

SELECT N'16(2) 20312班选修1号课程且不及格的学生视图 Stu_20312_2' AS Item;
SELECT Sno, Sname, Cno, Gmark FROM Stu_20312_2 ORDER BY Sno;

SELECT N'16(3) 学生学号、姓名和年龄视图 Stu_age' AS Item;
SELECT * FROM Stu_age ORDER BY Sno;

SELECT N'16(4) 查询2000年以后出生的学生姓名' AS Item;
SELECT Sname
FROM Student
WHERE Sbirth >= '2000-01-01'
ORDER BY Sno;

SELECT N'16(5) 查询20312班选修1号课程且不及格学生的学号、姓名和年龄' AS Item;
SELECT v.Sno, v.Sname, a.Age
FROM Stu_20312_2 AS v
JOIN Stu_age AS a ON v.Sno = a.Sno
ORDER BY v.Sno;

SELECT N'16(6) 查询选课数超过两门学生的平均成绩和选课门数' AS Item;
SELECT s.Sno, s.Sname, AVG(g.Gmark) AS AvgMark, COUNT(*) AS CourseCount
FROM Student AS s
JOIN Grade AS g ON s.Sno = g.Sno
GROUP BY s.Sno, s.Sname
HAVING COUNT(*) > 2
ORDER BY AvgMark DESC;

SELECT N'16(7) 软件工程专业中比计算机科学与技术专业所有学生年龄小的学生' AS Item;
SELECT s.Sno, s.Sname, dbo.AgeOf(s.Sbirth) AS Age
FROM Student AS s
JOIN Class AS c ON s.Clno = c.Clno
WHERE c.Speciality = N'软件工程'
  AND dbo.AgeOf(s.Sbirth) < ALL (
      SELECT dbo.AgeOf(s2.Sbirth)
      FROM Student AS s2
      JOIN Class AS c2 ON s2.Clno = c2.Clno
      WHERE c2.Speciality = N'计算机科学与技术'
  )
ORDER BY s.Sno;

SELECT N'16(8) 每门课程平均成绩和不及格率' AS Item;
SELECT c.Cno, c.Cname,
       CAST(AVG(g.Gmark) AS DECIMAL(6,2)) AS AvgMark,
       SUM(CASE WHEN g.Gmark < 60 THEN 1 ELSE 0 END) AS FailCount,
       CAST(100.0 * SUM(CASE WHEN g.Gmark < 60 THEN 1 ELSE 0 END) / COUNT(*) AS DECIMAL(6,2)) AS FailPercent
FROM Course AS c
JOIN Grade AS g ON c.Cno = g.Cno
GROUP BY c.Cno, c.Cname
ORDER BY c.Cno;
GO

CREATE OR ALTER VIEW Class_grade
AS
SELECT s.Clno, g.Cno, c.Cname,
       CAST(AVG(g.Gmark) AS DECIMAL(6,2)) AS AvgMark
FROM Student AS s
JOIN Grade AS g ON s.Sno = g.Sno
JOIN Course AS c ON g.Cno = c.Cno
GROUP BY s.Clno, g.Cno, c.Cname;
GO

SELECT N'实验内容2：Class_grade视图反映每个班各选修课平均成绩' AS Item;
SELECT * FROM Class_grade ORDER BY Clno, Cno;

SELECT N'验证Class_grade聚合视图是否可以更新' AS Item;
BEGIN TRY
    EXEC(N'UPDATE Class_grade
           SET AvgMark = 90
           WHERE Clno = ''20311'' AND Cno = ''1'';');
END TRY
BEGIN CATCH
    SELECT ERROR_NUMBER() AS ErrorNumber, ERROR_MESSAGE() AS ErrorMessage;
END CATCH;
GO
"""


EXPERIMENTS = [
    ("setup", SETUP_SQL),
    ("experiment5", EXPERIMENT5_SQL),
    ("experiment6", EXPERIMENT6_SQL),
    ("experiment7", EXPERIMENT7_SQL),
]


def run(command: list[str], *, cwd: Path | None = None, check: bool = True) -> subprocess.CompletedProcess:
    completed = subprocess.run(
        command,
        cwd=str(cwd or ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        check=False,
    )
    if check and completed.returncode != 0:
        output = completed.stdout.decode("utf-8", errors="replace")
        raise RuntimeError(f"command failed ({completed.returncode}): {' '.join(command)}\n{output}")
    return completed


def write_sql_files() -> dict[str, Path]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path] = {}
    for name, sql in EXPERIMENTS:
        path = OUT_DIR / f"{name}.sql"
        path.write_text(sql.strip() + "\n", encoding="utf-8-sig")
        paths[name] = path
    return paths


def execute_sql_script(name: str, path: Path) -> str:
    container_path = f"/tmp/{path.name}"
    run(["docker", "cp", str(path), f"{CONTAINER}:{container_path}"])
    cmd = [
        "docker",
        "exec",
        CONTAINER,
        SQLCMD,
        "-S",
        "localhost",
        "-U",
        "sa",
        "-P",
        SA_PASSWORD,
        "-C",
        "-b",
        "-W",
        "-w",
        "220",
        "-i",
        container_path,
    ]
    completed = run(cmd, check=True)
    text = completed.stdout.decode("utf-8", errors="replace")
    out_path = OUT_DIR / f"{name}_output.txt"
    out_path.write_text(text, encoding="utf-8")
    return text


def find_font() -> str:
    candidates = [
        Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "simhei.ttf",
        Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "simsun.ttc",
        Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "msyh.ttc",
        Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "consola.ttf",
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return ""


def wrap_visual_line(line: str, font: ImageFont.FreeTypeFont, draw: ImageDraw.ImageDraw, width: int) -> list[str]:
    if not line:
        return [""]
    pieces: list[str] = []
    current = ""
    for ch in line:
        trial = current + ch
        bbox = draw.textbbox((0, 0), trial, font=font)
        if bbox[2] - bbox[0] <= width or not current:
            current = trial
        else:
            pieces.append(current)
            current = ch
    if current:
        pieces.append(current)
    return pieces


def output_to_images(name: str, title: str, text: str) -> list[Path]:
    image_dir = OUT_DIR / "screenshots"
    if name == "setup" and image_dir.exists():
        shutil.rmtree(image_dir)
    image_dir.mkdir(parents=True, exist_ok=True)
    font_path = find_font()
    body_font = ImageFont.truetype(font_path, 24) if font_path else ImageFont.load_default()
    title_font = ImageFont.truetype(font_path, 34) if font_path else ImageFont.load_default()
    temp = Image.new("RGB", (10, 10), "white")
    draw = ImageDraw.Draw(temp)

    raw_lines = text.replace("\r\n", "\n").splitlines()
    lines: list[str] = []
    for raw in raw_lines:
        lines.extend(wrap_visual_line(raw.expandtabs(4), body_font, draw, 1760))

    # Split by visual line count so each screenshot remains readable in Word.
    chunks = [lines[i : i + 48] for i in range(0, len(lines), 48)] or [[]]
    paths: list[Path] = []
    for idx, chunk in enumerate(chunks, start=1):
        width = 1900
        line_height = 34
        top = 92
        height = max(360, top + len(chunk) * line_height + 48)
        image = Image.new("RGB", (width, height), (250, 250, 250))
        draw = ImageDraw.Draw(image)
        draw.rectangle((0, 0, width, 64), fill=(35, 43, 54))
        draw.text((28, 14), title, fill="white", font=title_font)
        y = top
        for line in chunk:
            draw.text((28, y), line, fill=(20, 20, 20), font=body_font)
            y += line_height
        out = image_dir / f"{name}_{idx}.png"
        image.save(out)
        paths.append(out)
    return paths


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_paragraph_font(paragraph, font_name: str = "宋体", size: float = 12, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold


def add_body_paragraph(doc: Document, text: str = "", *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(14)
    r.bold = bold


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(18 if level == 1 else 14)
        run.bold = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(6)


def add_code_block(doc: Document, code: str, max_lines: int | None = None) -> None:
    lines = code.strip().splitlines()
    if max_lines is not None and len(lines) > max_lines:
        head = lines[: max_lines // 2]
        tail = lines[-max_lines // 2 :]
        lines = head + ["-- 中间部分见 SQL/generated 中的完整脚本文件。"] + tail
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F8")
    set_cell_margins(cell, top=100, bottom=100, start=120, end=120)
    cell.text = ""
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(10.5)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(8.5)


def add_picture_block(doc: Document, image_paths: list[Path], caption: str) -> None:
    doc.add_page_break()
    for index, image_path in enumerate(image_paths, start=1):
        if index > 1:
            doc.add_page_break()
        add_body_paragraph(doc, caption, bold=True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(6.3))


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(20)


def build_report(outputs: dict[str, str], images: dict[str, list[Path]], sql_paths: dict[str, Path]) -> None:
    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(110)
    r = p.add_run("《数据库系统原理》课程")
    r.font.name = "华文行楷"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "华文行楷")
    r.font.size = Pt(26)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("实验报告")
    r.font.name = "华文行楷"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "华文行楷")
    r.font.size = Pt(26)
    r.bold = True

    for _ in range(3):
        doc.add_paragraph()
    for line in [
        "主题名称：      第2次SQL实验报告",
        f"学    号：      {STUDENT_ID}",
        f"姓    名：      {STUDENT_NAME}",
        "实验环境：      Docker SQL Server 2022 Express（容器 sql2022-lab，localhost:14333）",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(20)
        r = p.add_run(line)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(14)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目  录")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(20)
    r.bold = True
    for line in ["上机实验五\t1", "上机实验六\t2", "上机实验七\t3"]:
        add_body_paragraph(doc, line)
    add_body_paragraph(doc, "【说明】")
    add_body_paragraph(doc, "1.根据上机实际要求完成情况填写；")
    add_body_paragraph(doc, "2.格式要求：中文字体：宋体小四；")
    add_body_paragraph(doc, "英文字体：Times New Roman，小四；")
    add_body_paragraph(doc, "行间距：20磅。")

    doc.add_page_break()
    add_heading(doc, "上机实验五", 1)
    add_heading(doc, "实验目的", 2)
    add_body_paragraph(doc, "掌握 SELECT 语句的嵌套使用方法，使用高级格式和完整格式完成 SQL Server 对表的复杂查询。")
    add_heading(doc, "实验环境", 2)
    add_body_paragraph(doc, "MS SQL Server：Docker 容器 sql2022-lab，镜像 mcr.microsoft.com/mssql/server:2022-latest，端口 localhost:14333。")
    add_body_paragraph(doc, "本次实验未使用 openGauss。")
    add_heading(doc, "实验报告内容", 2)
    add_body_paragraph(doc, "题目要求：针对习题3第10题建立的 GradeManager 四个表，验证习题3第13题和第14题的各项 SELECT 查询；回答嵌套查询、集合查询、连接查询、聚合函数使用位置等问题。")
    add_body_paragraph(doc, "SQL脚本如下（完整脚本文件：SQL/generated/experiment5.sql）：", bold=True)
    add_code_block(doc, sql_paths["experiment5"].read_text(encoding="utf-8-sig"), max_lines=220)
    add_picture_block(doc, images["experiment5"], "实验五 SQL Server 实际执行结果截图：")
    add_heading(doc, "实验中遇到的问题和总结", 2)
    add_body_paragraph(doc, "1. 使用 EXISTS 时，外层当前行只要能让子查询返回记录，WHERE 条件就为真；使用 NOT EXISTS 时，只有子查询不返回记录时条件才为真，适合表达“没有未选课程”等全称条件。")
    add_body_paragraph(doc, "2. UNION 会合并两个 SELECT 的结果并去重；UNION ALL 保留重复记录，通常不做去重排序，执行开销更低。")
    add_body_paragraph(doc, "3. 能使用连接查询也能使用嵌套查询时，通常优先选择连接查询，因为语义直观、优化器可选执行计划更多；涉及“存在/不存在”“至少包含全部”等逻辑时，嵌套查询更清晰。")
    add_body_paragraph(doc, "4. 聚合函数可以出现在 SELECT 目标列表、HAVING 子句和 ORDER BY 中；不能直接放在 WHERE 条件中，也不能作为 GROUP BY 的分组列名。")
    add_body_paragraph(doc, "总结：本实验练习了 IN、EXISTS、NOT EXISTS、GROUP BY、HAVING、窗口函数等查询形式，理解了简单查询与复杂查询在表达能力上的差异。")

    doc.add_page_break()
    add_heading(doc, "上机实验六", 1)
    add_heading(doc, "一、实验目的", 2)
    add_body_paragraph(doc, "掌握用交互式 SQL 语句对已建基本表进行存储操作，包括修改、删除、插入，并加深对数据完整性的理解。")
    add_heading(doc, "二、实验环境", 2)
    add_body_paragraph(doc, "MS SQL Server：Docker 容器 sql2022-lab，镜像 mcr.microsoft.com/mssql/server:2022-latest，端口 localhost:14333。")
    add_body_paragraph(doc, "本次实验未使用 openGauss。")
    add_heading(doc, "三、实验报告内容", 2)
    add_body_paragraph(doc, "题目要求：针对习题3第10题建立的四个表，验证习题3第15题的 UPDATE、DELETE、SELECT INTO 等存储操作；判断工程师 Basepay 增加 100 的 UPDATE 语句是否正确；说明 DROP 与 DELETE 的区别。")
    add_body_paragraph(doc, "SQL脚本如下（完整脚本文件：SQL/generated/experiment6.sql）：", bold=True)
    add_code_block(doc, sql_paths["experiment6"].read_text(encoding="utf-8-sig"), max_lines=180)
    add_picture_block(doc, images["experiment6"], "实验六 SQL Server 实际执行结果截图：")
    add_heading(doc, "四、实验中遇到的问题和总结", 2)
    add_body_paragraph(doc, "1. 工程师基本工资增加 100 的 UPDATE 语句是正确的，条件子查询返回工程师的 Eno，外层 Salary 表按这些 Eno 更新 Basepay。SQL Server 中中文常量建议写成 N'工程师'。")
    add_body_paragraph(doc, "2. DROP 删除的是数据库对象本身，例如表结构、数据及相关对象定义都会被移除；DELETE 删除的是表中的行，表结构仍然保留，并且可以通过 WHERE 精确限定删除范围。")
    add_body_paragraph(doc, "3. 为避免实验六的更新和删除影响后续实验，本报告把破坏性操作放入事务内演示，显示执行结果后 ROLLBACK 恢复初始数据。")
    add_body_paragraph(doc, "总结：本实验验证了数据更新、删除和结果入库操作，理解了数据操作语句与数据库对象定义语句的边界。")

    doc.add_page_break()
    add_heading(doc, "上机实验七", 1)
    add_heading(doc, "实验目的", 2)
    add_body_paragraph(doc, "掌握创建、删除和查询视图的方法，验证可更新视图和不可更新视图的差异。")
    add_heading(doc, "实验环境", 2)
    add_body_paragraph(doc, "MS SQL Server：Docker 容器 sql2022-lab，镜像 mcr.microsoft.com/mssql/server:2022-latest，端口 localhost:14333。")
    add_body_paragraph(doc, "本次实验未使用 openGauss。")
    add_heading(doc, "实验报告内容", 2)
    add_body_paragraph(doc, "题目要求：完成习题3第16题的视图创建与查询；建立 Class_grade 视图反映每个班各选修课平均成绩，并验证该视图能否更新。")
    add_body_paragraph(doc, "SQL脚本如下（完整脚本文件：SQL/generated/experiment7.sql）：", bold=True)
    add_code_block(doc, sql_paths["experiment7"].read_text(encoding="utf-8-sig"), max_lines=200)
    add_picture_block(doc, images["experiment7"], "实验七 SQL Server 实际执行结果截图：")
    add_heading(doc, "实验中遇到的问题和总结", 2)
    add_body_paragraph(doc, "1. Stu_20312_1、Stu_20312_2 和 Stu_age 都能正常查询，它们把常用查询条件固化为逻辑表，简化了后续访问。")
    add_body_paragraph(doc, "2. Class_grade 视图包含 GROUP BY 和 AVG 聚合结果，SQL Server 不允许直接更新这种聚合视图，因为一行视图结果对应多行基础表记录，系统无法唯一确定应修改哪些基础数据。")
    add_body_paragraph(doc, "总结：视图可以提高查询复用性和安全性，但是否可更新取决于视图定义是否能明确映射到底层基本表。")

    doc.save(REPORT_PATH)


def export_pdf_and_render() -> list[Path]:
    import win32com.client as win32

    if RENDER_DIR.exists():
        shutil.rmtree(RENDER_DIR)
    RENDER_DIR.mkdir(parents=True, exist_ok=True)
    word = win32.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(str(REPORT_PATH.resolve()), ReadOnly=True, AddToRecentFiles=False)
        doc.ExportAsFixedFormat(str(PDF_PATH.resolve()), 17)
        doc.Close(False)
    finally:
        try:
            word.Quit()
        except Exception:
            pass

    pdf = pdfium.PdfDocument(str(PDF_PATH))
    rendered: list[Path] = []
    for i in range(len(pdf)):
        page = pdf[i]
        image = page.render(scale=1.4).to_pil().convert("RGB")
        out = RENDER_DIR / f"page-{i + 1:02d}.png"
        image.save(out)
        rendered.append(out)
    return rendered


def build_contact_sheets(rendered: list[Path]) -> list[Path]:
    sheets: list[Path] = []
    cols = 3
    thumb_w, thumb_h = 360, 520
    for start in range(0, len(rendered), 9):
        batch = rendered[start : start + 9]
        rows = (len(batch) + cols - 1) // cols
        sheet = Image.new("RGB", (cols * thumb_w, rows * thumb_h), (242, 242, 242))
        font_path = find_font()
        font = ImageFont.truetype(font_path, 18) if font_path else ImageFont.load_default()
        draw = ImageDraw.Draw(sheet)
        for idx, path in enumerate(batch):
            img = Image.open(path).convert("RGB")
            img.thumbnail((thumb_w - 20, thumb_h - 40))
            x = (idx % cols) * thumb_w
            y = (idx // cols) * thumb_h
            draw.text((x + 10, y + 8), f"Page {start + idx + 1}", fill=(0, 0, 0), font=font)
            sheet.paste(img, (x + (thumb_w - img.width) // 2, y + 34))
        out = RENDER_DIR / f"contact-{start // 9 + 1}.jpg"
        sheet.save(out, quality=90)
        sheets.append(out)
    return sheets


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    sql_paths = write_sql_files()

    outputs: dict[str, str] = {}
    for name in ["setup", "experiment5", "experiment6", "experiment7"]:
        outputs[name] = execute_sql_script(name, sql_paths[name])

    images = {
        "setup": output_to_images("setup", "公共数据库初始化执行结果", outputs["setup"]),
        "experiment5": output_to_images("experiment5", "上机实验五执行结果", outputs["experiment5"]),
        "experiment6": output_to_images("experiment6", "上机实验六执行结果", outputs["experiment6"]),
        "experiment7": output_to_images("experiment7", "上机实验七执行结果", outputs["experiment7"]),
    }

    build_report(outputs, images, sql_paths)
    rendered = export_pdf_and_render()
    sheets = build_contact_sheets(rendered)

    print(f"REPORT={REPORT_PATH}")
    print(f"PDF={PDF_PATH}")
    print(f"PAGES={len(rendered)}")
    for sheet in sheets:
        print(f"CONTACT={sheet}")


if __name__ == "__main__":
    main()
