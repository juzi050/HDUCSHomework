from __future__ import annotations

import os
import shutil
import subprocess
from pathlib import Path

import pypdfium2 as pdfium
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
SQL_DIR = ROOT / "SQL"
OUT_DIR = SQL_DIR / "generated"
STUDENT_ID = os.environ.get("SQL_REPORT_STUDENT_ID", "<学号>")
STUDENT_NAME = os.environ.get("SQL_REPORT_STUDENT_NAME", "<姓名>")
REPORT_PATH = SQL_DIR / "《数据库系统原理》课程实验报告模板3_模板.docx"
PDF_PATH = OUT_DIR / "《数据库系统原理》课程实验报告模板3_模板.pdf"
RENDER_DIR = OUT_DIR / "rendered_pages_report3"
SCREENSHOT_DIR = OUT_DIR / "screenshots_report3"

CONTAINER = "sql2022-lab"
SA_PASSWORD = "CodexLab!2026"
SQLCMD = "/opt/mssql-tools18/bin/sqlcmd"


SETUP_SQL = r"""
USE master;
IF DB_ID(N'GradeManager') IS NOT NULL
BEGIN
    ALTER DATABASE GradeManager SET SINGLE_USER WITH ROLLBACK IMMEDIATE;
    DROP DATABASE GradeManager;
END;
GO

CREATE DATABASE GradeManager COLLATE Chinese_PRC_CI_AS;
GO

USE GradeManager;
GO

CREATE TABLE dbo.Class (
    Clno CHAR(5) NOT NULL,
    Speciality VARCHAR(20) NOT NULL,
    Inyear CHAR(4) NOT NULL,
    Number INTEGER NULL,
    Monitor CHAR(7) NULL,
    CONSTRAINT PK_Class PRIMARY KEY (Clno),
    CONSTRAINT CK_Class_Number CHECK (Number > 25 AND Number < 50)
);

CREATE TABLE dbo.Student (
    Sno CHAR(7) NOT NULL,
    Sname VARCHAR(20) NOT NULL,
    Ssex CHAR(2) NOT NULL CONSTRAINT DF_Student_Ssex DEFAULT N'男',
    Sbirth DATE NULL,
    Clno CHAR(5) NOT NULL,
    CONSTRAINT PK_Student PRIMARY KEY (Sno),
    CONSTRAINT CK_Student_Ssex CHECK (Ssex IN (N'男', N'女')),
    CONSTRAINT CK_Student_Sbirth CHECK (Sbirth IS NULL OR Sbirth < CONVERT(date, GETDATE())),
    CONSTRAINT FK_Student_Class FOREIGN KEY (Clno) REFERENCES dbo.Class(Clno) ON UPDATE CASCADE
);

CREATE TABLE dbo.Course (
    Cno CHAR(3) NOT NULL,
    Cname VARCHAR(20) NOT NULL,
    Ccredit SMALLINT NULL,
    Cpno CHAR(3) NULL,
    CONSTRAINT PK_Course PRIMARY KEY (Cno),
    CONSTRAINT CK_Course_Ccredit CHECK (Ccredit IS NULL OR Ccredit IN (1, 2, 3, 4, 5, 6)),
    CONSTRAINT FK_Course_PreCourse FOREIGN KEY (Cpno) REFERENCES dbo.Course(Cno)
);

CREATE TABLE dbo.Grade (
    Sno CHAR(7) NOT NULL,
    Cno CHAR(3) NOT NULL,
    Gmark NUMERIC(4,1) NULL,
    CONSTRAINT PK_Grade PRIMARY KEY (Sno, Cno),
    CONSTRAINT CK_Grade_Gmark CHECK (Gmark IS NULL OR (Gmark > 0 AND Gmark < 100)),
    CONSTRAINT FK_Grade_Student FOREIGN KEY (Sno) REFERENCES dbo.Student(Sno) ON DELETE CASCADE,
    CONSTRAINT FK_Grade_Course FOREIGN KEY (Cno) REFERENCES dbo.Course(Cno) ON UPDATE CASCADE
);

INSERT INTO dbo.Class (Clno, Speciality, Inyear, Number, Monitor) VALUES
('20311', N'软件工程', '2020', 35, NULL),
('20312', N'计算机科学与技术', '2020', 38, NULL),
('21311', N'软件工程', '2021', 40, NULL);

INSERT INTO dbo.Student (Sno, Sname, Ssex, Sbirth, Clno) VALUES
('2020101', N'李勇', N'男', '2002-08-09', '20311'),
('2020102', N'刘诗晨', N'女', '2003-04-01', '20311'),
('2020103', N'王一鸣', N'男', '2002-12-25', '20312'),
('2020104', N'张婷婷', N'女', '2002-10-01', '20312'),
('2021101', N'李勇敏', N'女', '2003-11-11', '21311'),
('2021102', N'贾向东', N'男', '2003-12-12', '21311'),
('2021103', N'陈宝玉', N'男', '2004-05-01', '21311'),
('2021104', N'张逸凡', N'男', '2005-01-01', '21311');

UPDATE dbo.Class SET Monitor = '2020101' WHERE Clno = '20311';
UPDATE dbo.Class SET Monitor = '2020103' WHERE Clno = '20312';
UPDATE dbo.Class SET Monitor = '2021103' WHERE Clno = '21311';

ALTER TABLE dbo.Class
ADD CONSTRAINT FK_Class_Monitor FOREIGN KEY (Monitor) REFERENCES dbo.Student(Sno);

INSERT INTO dbo.Course (Cno, Cname, Ccredit, Cpno) VALUES
('7', N'C语言', 4, NULL),
('3', N'数字电路设计', 2, NULL),
('8', N'计算机组成原理', 4, '3'),
('5', N'数据结构', 4, '7'),
('1', N'数据库系统原理', 4, '5'),
('2', N'计算机系统结构', 3, '8'),
('4', N'操作系统', 4, '8'),
('6', N'软件工程', 2, '1');

INSERT INTO dbo.Grade (Sno, Cno, Gmark) VALUES
('2020101', '1', 92), ('2020101', '3', 88), ('2020101', '5', 86),
('2020102', '1', 78), ('2020102', '6', 55),
('2020103', '3', 65), ('2020103', '6', 78), ('2020103', '5', 66),
('2020104', '1', 54), ('2020104', '6', 83),
('2021101', '2', 70), ('2021101', '4', 65),
('2021102', '2', 80), ('2021102', '4', 90),
('2020103', '1', 83), ('2020103', '2', 76), ('2020103', '4', 56), ('2020103', '7', 88);

SELECT N'GradeManager 实验8-10初始化完成' AS Info,
       (SELECT COUNT(*) FROM dbo.Student) AS StudentRows,
       (SELECT COUNT(*) FROM dbo.Course) AS CourseRows,
       (SELECT COUNT(*) FROM dbo.Class) AS ClassRows,
       (SELECT COUNT(*) FROM dbo.Grade) AS GradeRows;
GO
"""


EXPERIMENT8_SQL = r"""
USE GradeManager;
GO

PRINT N'【实验八】完整性约束的实现';

SELECT N'习题4第11题：四个表已加入主键、外键、CHECK、DEFAULT等完整性约束' AS Item;
SELECT tc.TABLE_NAME, tc.CONSTRAINT_TYPE, tc.CONSTRAINT_NAME
FROM INFORMATION_SCHEMA.TABLE_CONSTRAINTS AS tc
WHERE tc.TABLE_SCHEMA = 'dbo'
ORDER BY tc.TABLE_NAME, tc.CONSTRAINT_TYPE, tc.CONSTRAINT_NAME;

SELECT N'验证用户自定义完整性：课程学分只能为1到6' AS Item;
BEGIN TRY
    INSERT INTO dbo.Course (Cno, Cname, Ccredit, Cpno)
    VALUES ('999', N'错误学分课程', 8, NULL);
END TRY
BEGIN CATCH
    SELECT ERROR_NUMBER() AS ErrorNumber, ERROR_MESSAGE() AS ErrorMessage;
END CATCH;
GO

CREATE OR ALTER TRIGGER dbo.trg_Student_ClassNumberAndLimit
ON dbo.Student
AFTER INSERT, DELETE, UPDATE
AS
BEGIN
    SET NOCOUNT ON;

    DECLARE @Delta TABLE (
        Clno CHAR(5) NOT NULL PRIMARY KEY,
        DeltaCount INT NOT NULL
    );

    INSERT INTO @Delta (Clno, DeltaCount)
    SELECT Clno, SUM(DeltaCount)
    FROM (
        SELECT Clno, COUNT(*) AS DeltaCount FROM inserted GROUP BY Clno
        UNION ALL
        SELECT Clno, -COUNT(*) AS DeltaCount FROM deleted GROUP BY Clno
    ) AS x
    GROUP BY Clno
    HAVING SUM(DeltaCount) <> 0;

    IF EXISTS (
        SELECT 1
        FROM dbo.Class AS c
        JOIN @Delta AS d ON c.Clno = d.Clno
        WHERE ISNULL(c.Number, 0) + d.DeltaCount > 40
    )
    BEGIN
        RAISERROR(N'该班学生人数超过40人，操作回滚。', 16, 1);
        ROLLBACK TRANSACTION;
        RETURN;
    END;

    UPDATE c
    SET Number = ISNULL(c.Number, 0) + d.DeltaCount
    FROM dbo.Class AS c
    JOIN @Delta AS d ON c.Clno = d.Clno;
END;
GO

CREATE OR ALTER TRIGGER dbo.trg_Class_MonitorCheck
ON dbo.Class
AFTER UPDATE
AS
BEGIN
    SET NOCOUNT ON;

    IF UPDATE(Monitor)
       AND EXISTS (
           SELECT 1
           FROM inserted AS i
           LEFT JOIN dbo.Student AS s ON i.Monitor = s.Sno
           WHERE i.Monitor IS NOT NULL
             AND (s.Sno IS NULL OR s.Clno <> i.Clno)
       )
    BEGIN
        RAISERROR(N'班长学号必须属于本班学生，更新取消。', 16, 1);
        ROLLBACK TRANSACTION;
        RETURN;
    END;
END;
GO

DROP VIEW IF EXISTS dbo.vw_StudentCredit;
DROP TABLE IF EXISTS dbo.StudentCredit;
GO

CREATE TABLE dbo.StudentCredit (
    Sno CHAR(7) NOT NULL PRIMARY KEY,
    EarnedCredit INT NOT NULL CONSTRAINT DF_StudentCredit_EarnedCredit DEFAULT 0,
    CONSTRAINT FK_StudentCredit_Student FOREIGN KEY (Sno) REFERENCES dbo.Student(Sno) ON DELETE CASCADE
);

INSERT INTO dbo.StudentCredit (Sno, EarnedCredit)
SELECT s.Sno,
       COALESCE(SUM(CASE WHEN g.Gmark >= 60 THEN c.Ccredit ELSE 0 END), 0) AS EarnedCredit
FROM dbo.Student AS s
LEFT JOIN dbo.Grade AS g ON s.Sno = g.Sno
LEFT JOIN dbo.Course AS c ON g.Cno = c.Cno
GROUP BY s.Sno;
GO

CREATE OR ALTER TRIGGER dbo.trg_Grade_UpdateStudentCredit
ON dbo.Grade
AFTER INSERT, DELETE, UPDATE
AS
BEGIN
    SET NOCOUNT ON;

    WITH Affected AS (
        SELECT Sno FROM inserted
        UNION
        SELECT Sno FROM deleted
    ),
    Calc AS (
        SELECT a.Sno,
               COALESCE(SUM(CASE WHEN g.Gmark >= 60 THEN c.Ccredit ELSE 0 END), 0) AS EarnedCredit
        FROM Affected AS a
        LEFT JOIN dbo.Grade AS g ON a.Sno = g.Sno
        LEFT JOIN dbo.Course AS c ON g.Cno = c.Cno
        GROUP BY a.Sno
    )
    MERGE dbo.StudentCredit AS target
    USING Calc AS source
       ON target.Sno = source.Sno
    WHEN MATCHED THEN
        UPDATE SET EarnedCredit = source.EarnedCredit
    WHEN NOT MATCHED THEN
        INSERT (Sno, EarnedCredit) VALUES (source.Sno, source.EarnedCredit);
END;
GO

CREATE OR ALTER VIEW dbo.vw_StudentCredit
AS
SELECT s.Sno, s.Sname, sc.EarnedCredit
FROM dbo.Student AS s
JOIN dbo.StudentCredit AS sc ON s.Sno = sc.Sno;
GO

SELECT N'习题4第12题：插入/删除Student时，Class.Number自动维护' AS Item;
SELECT N'插入前' AS Stage, Clno, Number FROM dbo.Class WHERE Clno = '20311';
INSERT INTO dbo.Student (Sno, Sname, Ssex, Sbirth, Clno)
VALUES ('2023999', N'测试学生', N'男', '2004-01-01', '20311');
SELECT N'插入后' AS Stage, Clno, Number FROM dbo.Class WHERE Clno = '20311';
DELETE FROM dbo.Student WHERE Sno = '2023999';
SELECT N'删除后' AS Stage, Clno, Number FROM dbo.Class WHERE Clno = '20311';

SELECT N'习题4第15题：班级人数超过40人时回滚' AS Item;
BEGIN TRY
    INSERT INTO dbo.Student (Sno, Sname, Ssex, Sbirth, Clno)
    VALUES ('2023998', N'超员测试', N'女', '2004-02-01', '21311');
END TRY
BEGIN CATCH
    SELECT ERROR_NUMBER() AS ErrorNumber, ERROR_MESSAGE() AS ErrorMessage;
END CATCH;
SELECT Clno, Number FROM dbo.Class WHERE Clno = '21311';

SELECT N'习题4第13题：班长必须为本班学生' AS Item;
BEGIN TRY
    UPDATE dbo.Class SET Monitor = '2020103' WHERE Clno = '20311';
END TRY
BEGIN CATCH
    SELECT ERROR_NUMBER() AS ErrorNumber, ERROR_MESSAGE() AS ErrorMessage;
END CATCH;
UPDATE dbo.Class SET Monitor = '2020102' WHERE Clno = '20311';
SELECT Clno, Monitor FROM dbo.Class WHERE Clno = '20311';

SELECT N'习题4第14题：成绩达到60分后才获得该课程学分' AS Item;
SELECT N'修改前' AS Stage, * FROM dbo.vw_StudentCredit WHERE Sno = '2020104';
BEGIN TRAN;
UPDATE dbo.Grade SET Gmark = 61 WHERE Sno = '2020104' AND Cno = '1';
SELECT N'事务内修改后' AS Stage, * FROM dbo.vw_StudentCredit WHERE Sno = '2020104';
ROLLBACK;
SELECT N'回滚后' AS Stage, * FROM dbo.vw_StudentCredit WHERE Sno = '2020104';
GO
"""


EXPERIMENT9_SQL = r"""
USE GradeManager;
GO

PRINT N'【实验九】安全性的实现';

DROP VIEW IF EXISTS dbo.vw_CS_CourseStats;
GO

IF DATABASE_PRINCIPAL_ID(N'计算机专业负责人') IS NOT NULL DROP ROLE [计算机专业负责人];
IF DATABASE_PRINCIPAL_ID(N'张老师') IS NOT NULL DROP USER [张老师];
IF DATABASE_PRINCIPAL_ID(N'李老师') IS NOT NULL DROP USER [李老师];
IF DATABASE_PRINCIPAL_ID(N'张勇') IS NOT NULL DROP USER [张勇];
IF DATABASE_PRINCIPAL_ID(N'李勇') IS NOT NULL DROP USER [李勇];
IF DATABASE_PRINCIPAL_ID(N'李勇敏') IS NOT NULL DROP USER [李勇敏];
GO

CREATE USER [张勇] WITHOUT LOGIN;
CREATE USER [李勇] WITHOUT LOGIN;
CREATE USER [李勇敏] WITHOUT LOGIN;
CREATE USER [张老师] WITHOUT LOGIN;
CREATE USER [李老师] WITHOUT LOGIN;
GO

SELECT N'习题4第17题(1)：张勇获得Student、Course查询权限' AS Item;
GRANT SELECT ON dbo.Student TO [张勇];
GRANT SELECT ON dbo.Course TO [张勇];

SELECT N'习题4第17题(2)：张勇获得Student插入、删除权限且可再授权' AS Item;
GRANT INSERT, DELETE ON dbo.Student TO [张勇] WITH GRANT OPTION;
EXECUTE AS USER = N'张勇';
GRANT INSERT, DELETE ON dbo.Student TO [李勇];
REVERT;

SELECT N'习题4第17题(3)：李勇获得Course查询权限和Ccredit修改权限' AS Item;
GRANT SELECT ON dbo.Course TO [李勇];
GRANT UPDATE (Ccredit) ON dbo.Course TO [李勇];

SELECT N'习题4第17题(4)：李勇敏获得Student全部数据操作权限且可再授权' AS Item;
GRANT SELECT, INSERT, UPDATE, DELETE ON dbo.Student TO [李勇敏] WITH GRANT OPTION;
EXECUTE AS USER = N'李勇敏';
GRANT SELECT ON dbo.Student TO [李勇];
REVERT;

SELECT N'授权结果查询' AS Item;
SELECT USER_NAME(dp.grantee_principal_id) AS Grantee,
       dp.permission_name,
       dp.state_desc,
       OBJECT_NAME(dp.major_id) AS ObjectName,
       COL_NAME(dp.major_id, dp.minor_id) AS ColumnName
FROM sys.database_permissions AS dp
WHERE USER_NAME(dp.grantee_principal_id) IN (N'张勇', N'李勇', N'李勇敏')
ORDER BY Grantee, ObjectName, permission_name, ColumnName;

SELECT N'习题4第17题(5)：撤销张勇在Student、Course上的查询权限' AS Item;
REVOKE SELECT ON dbo.Student FROM [张勇];
REVOKE SELECT ON dbo.Course FROM [张勇];

SELECT N'习题4第17题(6)：撤销张勇Student插入、删除权限，并级联撤销其再授权结果' AS Item;
REVOKE INSERT, DELETE ON dbo.Student FROM [张勇] CASCADE;

SELECT N'撤销后结果查询' AS Item;
SELECT USER_NAME(dp.grantee_principal_id) AS Grantee,
       dp.permission_name,
       dp.state_desc,
       OBJECT_NAME(dp.major_id) AS ObjectName,
       COL_NAME(dp.major_id, dp.minor_id) AS ColumnName
FROM sys.database_permissions AS dp
WHERE USER_NAME(dp.grantee_principal_id) IN (N'张勇', N'李勇', N'李勇敏')
ORDER BY Grantee, ObjectName, permission_name, ColumnName;

SELECT N'习题4第17题(7)：张勇、李勇获得建表和建存储过程权限' AS Item;
GRANT CREATE TABLE, CREATE PROCEDURE TO [张勇], [李勇];

GO

CREATE OR ALTER VIEW dbo.vw_CS_CourseStats
AS
SELECT c.Cno, c.Cname,
       COUNT(g.Sno) AS SelectCount,
       CAST(AVG(g.Gmark) AS DECIMAL(6,2)) AS AvgMark,
       MAX(g.Gmark) AS MaxMark,
       MIN(g.Gmark) AS MinMark
FROM dbo.Class AS cl
JOIN dbo.Student AS s ON cl.Clno = s.Clno
JOIN dbo.Grade AS g ON s.Sno = g.Sno
JOIN dbo.Course AS c ON g.Cno = c.Cno
WHERE cl.Speciality = N'计算机科学与技术'
GROUP BY c.Cno, c.Cname;
GO

CREATE ROLE [计算机专业负责人];
ALTER ROLE [计算机专业负责人] ADD MEMBER [张老师];
ALTER ROLE [计算机专业负责人] ADD MEMBER [李老师];
GRANT SELECT ON dbo.vw_CS_CourseStats TO [计算机专业负责人];
GO

SELECT N'习题4第18题：角色只能查看计算机科学与技术专业课程成绩统计视图' AS Item;
EXECUTE AS USER = N'张老师';
SELECT * FROM dbo.vw_CS_CourseStats ORDER BY Cno;
BEGIN TRY
    EXEC(N'SELECT TOP (1) Sno, Sname FROM dbo.Student;');
END TRY
BEGIN CATCH
    SELECT ERROR_NUMBER() AS ErrorNumber, ERROR_MESSAGE() AS ErrorMessage;
END CATCH;
REVERT;

SELECT N'角色成员与授权结果' AS Item;
SELECT rp.name AS RoleName, mp.name AS MemberName
FROM sys.database_role_members AS drm
JOIN sys.database_principals AS rp ON drm.role_principal_id = rp.principal_id
JOIN sys.database_principals AS mp ON drm.member_principal_id = mp.principal_id
WHERE rp.name = N'计算机专业负责人'
ORDER BY MemberName;
GO
"""


EXPERIMENT10_SQL = r"""
USE GradeManager;
GO

PRINT N'【实验十】创建存储过程和用户自定义函数';
GO

CREATE OR ALTER PROCEDURE dbo.usp_QueryStudentGrades
    @Sno CHAR(7)
AS
BEGIN
    SET NOCOUNT ON;
    SELECT s.Sno, s.Sname, c.Cno, c.Cname, g.Gmark
    FROM dbo.Student AS s
    JOIN dbo.Grade AS g ON s.Sno = g.Sno
    JOIN dbo.Course AS c ON g.Cno = c.Cno
    WHERE s.Sno = @Sno
    ORDER BY c.Cno;
END;
GO

CREATE OR ALTER PROCEDURE dbo.usp_ClassHasStudents
    @Clno CHAR(5),
    @ExistsFlag INT OUTPUT
AS
BEGIN
    SET NOCOUNT ON;
    IF EXISTS (SELECT 1 FROM dbo.Student WHERE Clno = @Clno)
        SET @ExistsFlag = 1;
    ELSE
        SET @ExistsFlag = 0;
END;
GO

CREATE OR ALTER PROCEDURE dbo.usp_QueryStudentNameAndMajor
    @Sno CHAR(7),
    @Sname VARCHAR(20) OUTPUT,
    @Speciality VARCHAR(20) OUTPUT
AS
BEGIN
    SET NOCOUNT ON;
    SELECT @Sname = s.Sname,
           @Speciality = c.Speciality
    FROM dbo.Student AS s
    JOIN dbo.Class AS c ON s.Clno = c.Clno
    WHERE s.Sno = @Sno;
END;
GO

CREATE OR ALTER FUNCTION dbo.fn_StudentGrades(@Sno CHAR(7))
RETURNS TABLE
AS
RETURN
(
    SELECT s.Sno, s.Sname, c.Cno, c.Cname, g.Gmark
    FROM dbo.Student AS s
    JOIN dbo.Grade AS g ON s.Sno = g.Sno
    JOIN dbo.Course AS c ON g.Cno = c.Cno
    WHERE s.Sno = @Sno
);
GO

CREATE OR ALTER FUNCTION dbo.fn_ClassHasStudents(@Clno CHAR(5))
RETURNS INT
AS
BEGIN
    DECLARE @ExistsFlag INT;
    IF EXISTS (SELECT 1 FROM dbo.Student WHERE Clno = @Clno)
        SET @ExistsFlag = 1;
    ELSE
        SET @ExistsFlag = 0;
    RETURN @ExistsFlag;
END;
GO

CREATE OR ALTER FUNCTION dbo.fn_StudentNameAndMajor(@Sno CHAR(7))
RETURNS TABLE
AS
RETURN
(
    SELECT s.Sno, s.Sname, c.Speciality
    FROM dbo.Student AS s
    JOIN dbo.Class AS c ON s.Clno = c.Clno
    WHERE s.Sno = @Sno
);
GO

SELECT N'习题5第17题：按学号查询该学生所有选修课成绩的存储过程' AS Item;
EXEC dbo.usp_QueryStudentGrades @Sno = '2020103';

SELECT N'习题5第18题：按班级号判断该班是否有学生的存储过程' AS Item;
DECLARE @HasStudent INT;
EXEC dbo.usp_ClassHasStudents @Clno = '20312', @ExistsFlag = @HasStudent OUTPUT;
SELECT '20312' AS Clno, @HasStudent AS HasStudent;
EXEC dbo.usp_ClassHasStudents @Clno = '99999', @ExistsFlag = @HasStudent OUTPUT;
SELECT '99999' AS Clno, @HasStudent AS HasStudent;

SELECT N'习题5第19题：按学号输出姓名和专业的存储过程' AS Item;
DECLARE @Name VARCHAR(20), @Major VARCHAR(20);
EXEC dbo.usp_QueryStudentNameAndMajor
    @Sno = '2020103',
    @Sname = @Name OUTPUT,
    @Speciality = @Major OUTPUT;
SELECT '2020103' AS Sno, @Name AS Sname, @Major AS Speciality;

SELECT N'习题5第20题：将第17-19题改写为用户自定义函数' AS Item;
SELECT * FROM dbo.fn_StudentGrades('2020103') ORDER BY Cno;
SELECT '20312' AS Clno, dbo.fn_ClassHasStudents('20312') AS HasStudent,
       '99999' AS EmptyClno, dbo.fn_ClassHasStudents('99999') AS EmptyClassHasStudent;
SELECT * FROM dbo.fn_StudentNameAndMajor('2020103');
GO
"""


EXPERIMENTS = [
    ("setup_report3", SETUP_SQL),
    ("experiment8", EXPERIMENT8_SQL),
    ("experiment9", EXPERIMENT9_SQL),
    ("experiment10", EXPERIMENT10_SQL),
]


def run(command: list[str], *, cwd: Path | None = None, check: bool = True) -> subprocess.CompletedProcess:
    completed = subprocess.run(
        command,
        cwd=str(cwd or ROOT),
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        check=False,
    )
    if check and completed.returncode != 0:
        output = completed.stdout.decode("utf-8", errors="replace")
        raise RuntimeError(f"command failed ({completed.returncode}): {' '.join(command)}\n{output}")
    return completed


def write_sql_files() -> dict[str, Path]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path] = {}
    for name, sql in EXPERIMENTS:
        path = OUT_DIR / f"{name}.sql"
        path.write_text(sql.strip() + "\n", encoding="utf-8-sig")
        paths[name] = path
    return paths


def execute_sql_script(name: str, path: Path) -> str:
    container_path = f"/tmp/{path.name}"
    run(["docker", "cp", str(path), f"{CONTAINER}:{container_path}"])
    cmd = [
        "docker",
        "exec",
        CONTAINER,
        SQLCMD,
        "-S",
        "localhost",
        "-U",
        "sa",
        "-P",
        SA_PASSWORD,
        "-C",
        "-b",
        "-W",
        "-w",
        "220",
        "-i",
        container_path,
    ]
    completed = run(cmd, check=True)
    text = completed.stdout.decode("utf-8", errors="replace")
    out_path = OUT_DIR / f"{name}_output.txt"
    out_path.write_text(text, encoding="utf-8")
    return text


def find_font() -> str:
    candidates = [
        Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "simhei.ttf",
        Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "simsun.ttc",
        Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "msyh.ttc",
        Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "consola.ttf",
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return ""


def wrap_visual_line(line: str, font: ImageFont.FreeTypeFont, draw: ImageDraw.ImageDraw, width: int) -> list[str]:
    if not line:
        return [""]
    pieces: list[str] = []
    current = ""
    for ch in line:
        trial = current + ch
        bbox = draw.textbbox((0, 0), trial, font=font)
        if bbox[2] - bbox[0] <= width or not current:
            current = trial
        else:
            pieces.append(current)
            current = ch
    if current:
        pieces.append(current)
    return pieces


def output_to_images(name: str, title: str, text: str) -> list[Path]:
    if name == "setup_report3" and SCREENSHOT_DIR.exists():
        shutil.rmtree(SCREENSHOT_DIR)
    SCREENSHOT_DIR.mkdir(parents=True, exist_ok=True)
    font_path = find_font()
    body_font = ImageFont.truetype(font_path, 24) if font_path else ImageFont.load_default()
    title_font = ImageFont.truetype(font_path, 34) if font_path else ImageFont.load_default()
    temp = Image.new("RGB", (10, 10), "white")
    draw = ImageDraw.Draw(temp)

    raw_lines = text.replace("\r\n", "\n").splitlines()
    lines: list[str] = []
    for raw in raw_lines:
        lines.extend(wrap_visual_line(raw.expandtabs(4), body_font, draw, 1760))

    chunks = [lines[i : i + 48] for i in range(0, len(lines), 48)] or [[]]
    paths: list[Path] = []
    for idx, chunk in enumerate(chunks, start=1):
        width = 1900
        line_height = 34
        top = 92
        height = max(360, top + len(chunk) * line_height + 48)
        image = Image.new("RGB", (width, height), (250, 250, 250))
        draw = ImageDraw.Draw(image)
        draw.rectangle((0, 0, width, 64), fill=(35, 43, 54))
        draw.text((28, 14), title, fill="white", font=title_font)
        y = top
        for line in chunk:
            draw.text((28, y), line, fill=(20, 20, 20), font=body_font)
            y += line_height
        out = SCREENSHOT_DIR / f"{name}_{idx}.png"
        image.save(out)
        paths.append(out)
    return paths


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(20)


def add_body_paragraph(doc: Document, text: str = "", *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(14)
    r.bold = bold


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(18 if level == 1 else 14)
        run.bold = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(6)


def add_code_block(doc: Document, code: str, max_lines: int | None = None) -> None:
    lines = code.strip().splitlines()
    if max_lines is not None and len(lines) > max_lines:
        head = lines[: max_lines // 2]
        tail = lines[-max_lines // 2 :]
        lines = head + ["-- 中间部分见 SQL/generated 中的完整脚本文件。"] + tail
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F8")
    set_cell_margins(cell, top=100, bottom=100, start=120, end=120)
    cell.text = ""
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(10.5)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(8.5)


def add_picture_block(doc: Document, image_paths: list[Path], caption: str) -> None:
    doc.add_page_break()
    for index, image_path in enumerate(image_paths, start=1):
        if index > 1:
            doc.add_page_break()
        add_body_paragraph(doc, caption, bold=True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(6.3))


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(110)
    r = p.add_run("《数据库系统原理》课程")
    r.font.name = "华文行楷"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "华文行楷")
    r.font.size = Pt(26)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("实验报告")
    r.font.name = "华文行楷"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "华文行楷")
    r.font.size = Pt(26)
    r.bold = True

    for _ in range(3):
        doc.add_paragraph()

    for line in [
        "主题名称：      第3次SQL实验报告",
        f"学    号：      {STUDENT_ID}",
        f"姓    名：      {STUDENT_NAME}",
        "实验环境：      Docker SQL Server 2022 Express（容器 sql2022-lab，localhost:14333）",
    ]:
        add_body_paragraph(doc, line)


def add_catalog(doc: Document) -> None:
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目  录")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(20)
    r.bold = True

    for line in ["上机实验八\t1", "上机实验九\t2", "上机实验十\t3"]:
        add_body_paragraph(doc, line)
    add_body_paragraph(doc, "【说明】")
    add_body_paragraph(doc, "1.根据上机实际要求完成情况填写；")
    add_body_paragraph(doc, "2.格式要求：中文字体：宋体小四；")
    add_body_paragraph(doc, "英文字体：Times New Roman，小四；")
    add_body_paragraph(doc, "行间距：20磅。")


def build_report(images: dict[str, list[Path]], sql_paths: dict[str, Path]) -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_catalog(doc)

    doc.add_page_break()
    add_heading(doc, "上机实验八", 1)
    add_heading(doc, "实验目的", 2)
    add_body_paragraph(doc, "掌握 SQL 中实现数据完整性的方法，加深理解实体完整性、参照完整性和用户自定义完整性约束；掌握触发器的概念、作用和创建方法。")
    add_heading(doc, "实验环境", 2)
    add_body_paragraph(doc, "MS SQL Server：Docker 容器 sql2022-lab，镜像 mcr.microsoft.com/mssql/server:2022-latest，端口 localhost:14333。")
    add_body_paragraph(doc, "本次实验未使用 openGauss。")
    add_heading(doc, "实验报告内容", 2)
    add_body_paragraph(doc, "题目要求：验证习题4第11题中加入完整性约束后的四个成绩管理表结构；创建并验证习题4第12题至第15题的触发器。")
    add_body_paragraph(doc, "SQL脚本如下（完整脚本文件：SQL/generated/experiment8.sql）：", bold=True)
    add_code_block(doc, sql_paths["experiment8"].read_text(encoding="utf-8-sig"), max_lines=220)
    add_picture_block(doc, images["experiment8"], "实验八 SQL Server 实际执行结果截图：")
    add_heading(doc, "实验中遇到的问题和总结", 2)
    add_body_paragraph(doc, "1. 主键、外键、CHECK 和 DEFAULT 约束可以直接写在建表语句中，由 DBMS 自动维护，适合表达稳定的数据合法性要求。")
    add_body_paragraph(doc, "2. 触发器适合处理跨表维护和较复杂的业务规则。本实验中，Student 表变化后自动维护 Class.Number，班级人数超过40人时回滚，班长学号不属于本班时取消更新。")
    add_body_paragraph(doc, "3. INSERTED 和 DELETED 是触发器中的临时逻辑表，分别保存插入/更新后的记录和删除/更新前的记录；批量操作时必须按集合处理，不能假设一次只影响一行。")
    add_body_paragraph(doc, "总结：完整性约束负责基础规则，触发器负责更复杂的自动维护和联动检查，两者配合可以提升数据一致性。")

    doc.add_page_break()
    add_heading(doc, "上机实验九", 1)
    add_heading(doc, "一、实验目的", 2)
    add_body_paragraph(doc, "理解 SQL Server 下的安全性机制，掌握数据库级和数据库对象级安全保护机制的设计与实现方法。")
    add_heading(doc, "二、实验环境", 2)
    add_body_paragraph(doc, "MS SQL Server：Docker 容器 sql2022-lab，镜像 mcr.microsoft.com/mssql/server:2022-latest，端口 localhost:14333。")
    add_body_paragraph(doc, "本次实验未使用 openGauss。")
    add_heading(doc, "三、实验报告内容", 2)
    add_body_paragraph(doc, "题目要求：完成习题4第17题、第18题的用户授权、权限撤销、角色创建和角色授权，并验证安全机制设置是否生效。")
    add_body_paragraph(doc, "SQL脚本如下（完整脚本文件：SQL/generated/experiment9.sql）：", bold=True)
    add_code_block(doc, sql_paths["experiment9"].read_text(encoding="utf-8-sig"), max_lines=220)
    add_picture_block(doc, images["experiment9"], "实验九 SQL Server 实际执行结果截图：")
    add_heading(doc, "四、实验中遇到的问题和总结", 2)
    add_body_paragraph(doc, "1. GRANT 用于授予权限，WITH GRANT OPTION 表示被授权用户还能继续把该权限授予他人；REVOKE 可撤销权限，必要时使用 CASCADE 级联撤销再授权结果。")
    add_body_paragraph(doc, "2. 用户是具体的数据库访问主体，角色是权限集合。把权限授予角色，再把用户加入角色，可以减少重复授权，便于统一维护。")
    add_body_paragraph(doc, "3. 为限制专业负责人只能查看本专业课程成绩统计，本实验创建了过滤后的视图，并把 SELECT 权限授予角色，避免直接开放基础表。")
    add_body_paragraph(doc, "总结：数据库安全控制的核心是最小权限原则，应优先通过视图、角色和对象级授权控制用户可访问的数据范围。")

    doc.add_page_break()
    add_heading(doc, "上机实验十", 1)
    add_heading(doc, "实验目的", 2)
    add_body_paragraph(doc, "掌握存储过程和用户自定义函数的概念、作用，掌握存储过程的定义和执行方法，以及用户自定义函数的定义和调用方法。")
    add_heading(doc, "实验环境", 2)
    add_body_paragraph(doc, "MS SQL Server：Docker 容器 sql2022-lab，镜像 mcr.microsoft.com/mssql/server:2022-latest，端口 localhost:14333。")
    add_body_paragraph(doc, "本次实验未使用 openGauss。")
    add_heading(doc, "实验报告内容", 2)
    add_body_paragraph(doc, "题目要求：创建并执行习题5第17题至第19题的存储过程；将第17题至第19题分别改写为用户自定义函数并调用验证；回答存储过程的作用、优点及其与函数的区别。")
    add_body_paragraph(doc, "SQL脚本如下（完整脚本文件：SQL/generated/experiment10.sql）：", bold=True)
    add_code_block(doc, sql_paths["experiment10"].read_text(encoding="utf-8-sig"), max_lines=220)
    add_picture_block(doc, images["experiment10"], "实验十 SQL Server 实际执行结果截图：")
    add_heading(doc, "实验中遇到的问题和总结", 2)
    add_body_paragraph(doc, "1. 存储过程把一组 SQL 操作封装在服务器端，可以接收输入参数和输出参数，便于复用业务逻辑、减少网络传输并统一管理数据库操作。")
    add_body_paragraph(doc, "2. 用户自定义函数更适合表达可被查询语句调用的计算或查询逻辑。表值函数可以像表一样在 SELECT 中使用，标量函数可以返回单个值。")
    add_body_paragraph(doc, "3. 存储过程可包含更完整的过程控制和数据修改逻辑，调用方式是 EXEC；函数通常要求有明确返回值，可在 SELECT、WHERE 等查询语句中调用，但受副作用限制更多。")
    add_body_paragraph(doc, "总结：存储过程偏向封装操作流程，函数偏向封装可复用的表达式或查询结果。根据是否需要输出参数、数据修改和查询内调用来选择。")

    doc.save(REPORT_PATH)


def export_pdf_and_render() -> list[Path]:
    import win32com.client as win32

    if RENDER_DIR.exists():
        shutil.rmtree(RENDER_DIR)
    RENDER_DIR.mkdir(parents=True, exist_ok=True)
    word = win32.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(str(REPORT_PATH.resolve()), ReadOnly=True, AddToRecentFiles=False)
        doc.ExportAsFixedFormat(str(PDF_PATH.resolve()), 17)
        doc.Close(False)
    finally:
        try:
            word.Quit()
        except Exception:
            pass

    pdf = pdfium.PdfDocument(str(PDF_PATH))
    rendered: list[Path] = []
    for i in range(len(pdf)):
        page = pdf[i]
        image = page.render(scale=1.4).to_pil().convert("RGB")
        out = RENDER_DIR / f"page-{i + 1:02d}.png"
        image.save(out)
        rendered.append(out)
    return rendered


def build_contact_sheets(rendered: list[Path]) -> list[Path]:
    sheets: list[Path] = []
    cols = 3
    thumb_w, thumb_h = 360, 520
    for start in range(0, len(rendered), 9):
        batch = rendered[start : start + 9]
        rows = (len(batch) + cols - 1) // cols
        sheet = Image.new("RGB", (cols * thumb_w, rows * thumb_h), (242, 242, 242))
        font_path = find_font()
        font = ImageFont.truetype(font_path, 18) if font_path else ImageFont.load_default()
        draw = ImageDraw.Draw(sheet)
        for idx, path in enumerate(batch):
            img = Image.open(path).convert("RGB")
            img.thumbnail((thumb_w - 20, thumb_h - 40))
            x = (idx % cols) * thumb_w
            y = (idx // cols) * thumb_h
            draw.text((x + 10, y + 8), f"Page {start + idx + 1}", fill=(0, 0, 0), font=font)
            sheet.paste(img, (x + (thumb_w - img.width) // 2, y + 34))
        out = RENDER_DIR / f"contact-{start // 9 + 1}.jpg"
        sheet.save(out, quality=90)
        sheets.append(out)
    return sheets


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    sql_paths = write_sql_files()

    outputs: dict[str, str] = {}
    for name in ["setup_report3", "experiment8", "experiment9", "experiment10"]:
        outputs[name] = execute_sql_script(name, sql_paths[name])

    images = {
        "experiment8": output_to_images("experiment8", "上机实验八执行结果", outputs["experiment8"]),
        "experiment9": output_to_images("experiment9", "上机实验九执行结果", outputs["experiment9"]),
        "experiment10": output_to_images("experiment10", "上机实验十执行结果", outputs["experiment10"]),
    }

    build_report(images, sql_paths)
    rendered = export_pdf_and_render()
    sheets = build_contact_sheets(rendered)

    print(f"REPORT={REPORT_PATH}")
    print(f"PDF={PDF_PATH}")
    print(f"PAGES={len(rendered)}")
    for sheet in sheets:
        print(f"CONTACT={sheet}")


if __name__ == "__main__":
    main()
